"""
KpopNara PO Generator — Cloud Function
Generates a Purchase Order document (docx) from approved invoice data.
Called by n8n after BigQuery write.
"""

import io
import json
import functions_framework
from docx import Document


@functions_framework.http
def generate_po(request):
    """
    POST with JSON: { po_id, po_number, vendor_name, line_items: [...], total_amount }
    Returns docx file as binary.
    """
    if request.method != "POST":
        return ("Method not allowed", 405)

    try:
        data = request.get_json(silent=True) or {}
        po_id = data.get("po_id", "N/A")
        po_number = data.get("po_number", "N/A")
        vendor_name = data.get("vendor_name", "N/A")
        line_items = data.get("line_items", [])
        total_amount = data.get("total_amount", 0)

        doc = Document()
        doc.add_heading("Purchase Order", 0)
        doc.add_paragraph(f"PO Number: {po_number}")
        doc.add_paragraph(f"PO ID: {po_id}")
        doc.add_paragraph(f"Vendor: {vendor_name}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "SKU"
        hdr[1].text = "Product"
        hdr[2].text = "Qty"
        hdr[3].text = "Unit Price"
        hdr[4].text = "Total"
        hdr[5].text = "Notes"

        for li in line_items:
            row = table.add_row().cells
            row[0].text = str(li.get("sku", li.get("matched_sku", "")))
            pn = li.get("product_name", li.get("matched_product_name", ""))
            vn = li.get("variant_name", "")
            display_name = f"{pn} ({vn})" if vn else pn
            row[1].text = str(display_name)
            row[2].text = str(li.get("quantity", ""))
            row[3].text = str(li.get("unit_price", ""))
            row[4].text = str(li.get("total_price", ""))
            row[5].text = str(li.get("vendor_notation", ""))

        doc.add_paragraph("")
        doc.add_paragraph(f"Total Amount: ${total_amount:,.2f}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return (
            buffer.getvalue(),
            200,
            {
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": f'attachment; filename="PO-{po_number}.docx"',
            },
        )
    except Exception as e:
        return (json.dumps({"error": str(e)}), 500, {"Content-Type": "application/json"})

